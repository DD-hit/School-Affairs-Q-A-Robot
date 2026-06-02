from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import glob
import os

def set_document_fonts(doc):
    """设置文档全局字体样式 - 标题仿宋，正文楷体小四"""
    # 设置正文样式为楷体小四
    style = doc.styles['Normal']
    style.font.name = '楷体'
    style.font.size = Pt(12)  # 小四号
    
    # 设置Heading样式为仿宋
    for level in [1, 2, 3]:
        style_name = f'Heading {level}'
        try:
            style = doc.styles[style_name]
            style.font.name = '仿宋'
            if level == 1:
                style.font.size = Pt(16)
            elif level == 2:
                style.font.size = Pt(14)
            else:
                style.font.size = Pt(12)
        except:
            pass
    
    # 手动设置Document的默认字体
    # 对于中文，需要设置 rFonts (East Asian Fonts)
    style = doc.styles['Normal']
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

def get_page_description(filename, soup):
    """获取页面的详细描述"""
    descriptions = {
        'admin.html': '管理员后台管理页面，提供系统管理功能，包括用户管理、内容审核、数据统计等。',
        'forum.html': '校务论坛主页面，展示论坛分类、热门帖子、最新讨论等内容。',
        'help.html': '系统使用帮助页面，提供FAQ、使用指南、常见问题解答等功能。',
        'index.html': '系统首页，提供智能问答入口、热门问题推荐、快速导航等功能。',
        'login.html': '用户登录和注册页面，提供账户认证功能。',
        'personal.html': '用户个人中心页面，管理个人信息、历史记录、收藏等内容。',
        'post.html': '帖子详情页面，展示帖子内容、评论、回复等功能。'
    }
    
    return descriptions.get(filename, '系统功能页面')

def infer_goal_and_outputs(soup, filename):
    """推断页面目标和输出"""
    title = soup.title.string.strip() if soup.title and soup.title.string else filename
    text = (soup.get_text('\n') or '').lower()
    
    # 根据页面内容推断目标和输出
    if 'admin' in filename or '管理员' in title or '后台' in title:
        goal = f"管理员后台管理 — {title}"
        outputs = [
            '展示系统统计数据（用户数、帖子数、问答数等）',
            '提供用户管理功能（查看、编辑、禁用用户）',
            '提供内容审核功能（审核帖子、评论）',
            '提供知识库管理功能（添加、编辑、删除问答对）',
            '提供系统设置功能（配置系统参数）'
        ]
    elif 'forum' in filename or '论坛' in title:
        goal = f"校务论坛浏览与交互 — {title}"
        outputs = [
            '展示论坛分类和帖子列表',
            '提供帖子搜索和筛选功能',
            '显示帖子统计信息（帖子数、回复数等）',
            '提供发帖和回复功能入口',
            '展示热门帖子和最新讨论'
        ]
    elif 'help' in filename or '帮助' in title:
        goal = f"系统帮助与支持 — {title}"
        outputs = [
            '提供分类帮助文档',
            '展示常见问题解答（FAQ）',
            '提供使用指南和教程',
            '提供联系支持方式',
            '展示系统功能介绍'
        ]
    elif 'index' in filename or '首页' in title or '问答' in text:
        goal = f"智能问答系统首页 — {title}"
        outputs = [
            '接收用户问题输入',
            '展示智能问答结果',
            '提供热门问题推荐',
            '展示系统功能导航',
            '提供快速搜索功能'
        ]
    elif 'login' in filename or '登录' in title:
        goal = f"用户认证与账户管理 — {title}"
        outputs = [
            '验证用户身份信息',
            '成功登录后跳转到首页或个人中心',
            '新用户注册功能',
            '密码找回功能',
            '记住登录状态功能'
        ]
    elif 'personal' in filename or '个人' in title:
        goal = f"用户个人中心管理 — {title}"
        outputs = [
            '展示用户个人信息',
            '管理历史问答记录',
            '管理收藏的帖子',
            '编辑个人资料',
            '查看消息通知'
        ]
    elif 'post' in filename or '帖子' in title:
        goal = f"帖子内容浏览与交互 — {title}"
        outputs = [
            '展示帖子详细内容',
            '显示评论和回复',
            '提供点赞/收藏功能',
            '提供评论/回复功能',
            '显示相关帖子推荐'
        ]
    else:
        goal = title
        outputs = ['展示页面内容或导航行为']
    
    return goal, outputs

def extract_inputs(soup):
    """提取页面输入元素"""
    inputs = []
    
    # 提取表单输入
    forms = soup.find_all('form')
    for form in forms:
        for inp in form.find_all(['input', 'select', 'textarea']):
            inp_type = inp.get('type', 'text')
            name = inp.get('name') or inp.get('id') or inp.get('placeholder') or inp_type
            required = '必填' if inp.get('required') else '可选'
            
            # 根据输入类型添加描述
            if inp_type == 'text':
                desc = f"文本输入框 - {name} ({required})"
            elif inp_type == 'password':
                desc = f"密码输入框 - {name} ({required})"
            elif inp_type == 'email':
                desc = f"邮箱输入框 - {name} ({required})"
            elif inp_type == 'checkbox':
                desc = f"复选框 - {name} ({required})"
            elif inp_type == 'radio':
                desc = f"单选按钮 - {name} ({required})"
            elif inp.name == 'select':
                desc = f"下拉选择框 - {name} ({required})"
            elif inp.name == 'textarea':
                desc = f"多行文本输入框 - {name} ({required})"
            else:
                desc = f"输入字段 - {name} ({required})"
            
            inputs.append(desc)
    
    # 提取搜索框
    search_inputs = soup.find_all(attrs={"class": lambda v: v and any(x in str(v).lower() for x in ['search', 'query'])})
    for si in search_inputs:
        if si.name in ['input', 'textarea']:
            placeholder = si.get('placeholder', '搜索')
            inputs.append(f"搜索输入框 - {placeholder}")
    
    # 提取按钮
    buttons = soup.find_all(['button', 'input'])
    for btn in buttons:
        if btn.get('type') in ['submit', 'button'] or btn.name == 'button':
            text = btn.get_text(strip=True) or btn.get('value') or '按钮'
            if text and len(text) < 50:  # 避免过长的按钮文本
                inputs.append(f"操作按钮 - {text}")
    
    return inputs

def get_use_cases(filename):
    """获取页面的用例文本"""
    use_cases = {
        'admin.html': {
            'participants': ['系统管理员'],
            'preconditions': ['管理员已登录系统', '拥有管理员权限'],
            'main_flow': [
                '管理员访问后台管理页面',
                '系统显示管理仪表板，包含统计数据',
                '管理员选择用户管理功能',
                '系统显示用户列表，包含搜索和筛选选项',
                '管理员查看或编辑用户信息',
                '系统保存更改并显示成功提示'
            ],
            'alternative_flow': [
                'A1：内容审核流程',
                '1. 管理员选择内容审核功能',
                '2. 系统显示待审核帖子列表',
                '3. 管理员审核帖子内容',
                '4. 管理员批准或拒绝帖子',
                '5. 系统更新帖子状态并通知发帖用户',
                '',
                'A2：知识库管理流程',
                '1. 管理员选择知识库管理功能',
                '2. 系统显示问答对列表',
                '3. 管理员添加新的问答对',
                '4. 管理员编辑或删除现有问答对',
                '5. 系统保存更改并更新知识库'
            ],
            'postconditions': ['管理操作已保存', '系统数据已更新', '相关用户收到通知（如需要）']
        },
        'forum.html': {
            'participants': ['普通用户', '注册用户'],
            'preconditions': ['用户已访问系统', '网络连接正常'],
            'main_flow': [
                '用户访问校务论坛页面',
                '系统显示论坛分类和帖子列表',
                '用户浏览帖子列表或使用搜索功能',
                '用户点击感兴趣的帖子',
                '系统跳转到帖子详情页面'
            ],
            'alternative_flow': [
                'A1：发帖流程（需登录）',
                '1. 用户点击"发帖"按钮',
                '2. 系统检查用户登录状态',
                '3. 如未登录，跳转到登录页面',
                '4. 如已登录，显示发帖表单',
                '5. 用户填写帖子标题和内容',
                '6. 用户选择帖子分类',
                '7. 用户提交帖子',
                '8. 系统保存帖子并显示成功提示',
                '',
                'A2：帖子筛选流程',
                '1. 用户选择分类标签进行筛选',
                '2. 系统根据选择筛选帖子列表',
                '3. 用户使用排序功能（按时间、热度等）',
                '4. 系统重新排序显示帖子'
            ],
            'postconditions': ['用户浏览了论坛内容', '可能产生了新的浏览记录', '如需发帖则进入发帖流程']
        },
        'help.html': {
            'participants': ['所有用户'],
            'preconditions': ['用户已访问系统'],
            'main_flow': [
                '用户访问帮助页面',
                '系统显示帮助分类导航',
                '用户选择感兴趣的帮助类别',
                '系统显示该类别下的帮助内容',
                '用户阅读帮助文档'
            ],
            'alternative_flow': [
                'A1：搜索帮助内容',
                '1. 用户在帮助页面使用搜索功能',
                '2. 系统根据关键词搜索帮助文档',
                '3. 系统显示搜索结果',
                '4. 用户点击查看相关帮助内容',
                '',
                'A2：联系支持',
                '1. 用户点击"联系支持"链接',
                '2. 系统显示联系方式或反馈表单',
                '3. 用户填写反馈信息',
                '4. 用户提交反馈',
                '5. 系统发送反馈确认'
            ],
            'postconditions': ['用户获得了所需帮助信息', '可能提交了反馈或问题']
        },
        'index.html': {
            'participants': ['所有用户'],
            'preconditions': ['用户已访问系统网站'],
            'main_flow': [
                '用户访问系统首页',
                '系统显示欢迎界面和主要功能',
                '用户在搜索框中输入问题',
                '用户点击搜索或按回车键',
                '系统处理问题并显示回答',
                '用户查看回答结果'
            ],
            'alternative_flow': [
                'A1：浏览热门问题',
                '1. 用户浏览热门问题推荐',
                '2. 用户点击感兴趣的问题',
                '3. 系统显示该问题的详细回答',
                '',
                'A2：使用功能导航',
                '1. 用户点击导航菜单中的功能链接',
                '2. 系统跳转到对应功能页面（如论坛、个人中心等）',
                '',
                'A3：未找到答案',
                '1. 系统未找到问题的直接答案',
                '2. 系统显示相关建议或引导到论坛提问',
                '3. 用户选择到论坛提问或重新表述问题'
            ],
            'postconditions': ['用户获得了问题答案', '可能产生了新的搜索记录', '可能跳转到其他功能页面']
        },
        'login.html': {
            'participants': ['未登录用户'],
            'preconditions': ['用户拥有系统账户（登录时）', '或准备注册新账户（注册时）'],
            'main_flow': [
                '用户访问登录页面',
                '用户输入用户名/邮箱和密码',
                '用户点击"登录"按钮',
                '系统验证用户凭证',
                '验证成功，系统跳转到首页或个人中心',
                '系统显示登录成功提示'
            ],
            'alternative_flow': [
                'A1：注册新账户',
                '1. 用户点击"注册"标签',
                '2. 系统显示注册表单',
                '3. 用户填写注册信息（用户名、邮箱、密码等）',
                '4. 用户同意服务条款',
                '5. 用户提交注册',
                '6. 系统创建新账户并自动登录',
                '7. 系统跳转到欢迎页面或首页',
                '',
                'A2：密码找回',
                '1. 用户点击"忘记密码"链接',
                '2. 系统显示密码找回表单',
                '3. 用户输入注册邮箱',
                '4. 系统发送密码重置链接到邮箱',
                '5. 用户通过邮件链接重置密码',
                '',
                'A3：登录失败',
                '1. 系统验证用户凭证失败',
                '2. 系统显示错误提示（用户名或密码错误）',
                '3. 用户重新输入或选择密码找回'
            ],
            'postconditions': ['用户成功登录系统', '或成功注册新账户', '系统记录登录状态']
        },
        'personal.html': {
            'participants': ['已登录用户'],
            'preconditions': ['用户已成功登录系统'],
            'main_flow': [
                '用户访问个人中心页面',
                '系统显示用户个人信息概览',
                '用户查看个人资料',
                '用户编辑需要修改的信息',
                '用户保存更改',
                '系统更新个人信息并显示成功提示'
            ],
            'alternative_flow': [
                'A1：查看历史记录',
                '1. 用户点击"历史记录"菜单',
                '2. 系统显示用户的问答历史',
                '3. 用户浏览或搜索历史记录',
                '4. 用户查看特定历史记录的详情',
                '',
                'A2：管理收藏',
                '1. 用户点击"我的收藏"菜单',
                '2. 系统显示用户收藏的帖子',
                '3. 用户管理收藏（取消收藏、查看详情）',
                '',
                'A3：消息通知',
                '1. 用户点击"消息通知"菜单',
                '2. 系统显示未读和已读消息',
                '3. 用户查看消息详情',
                '4. 用户标记消息为已读或删除消息'
            ],
            'postconditions': ['个人信息已更新', '用户查看了个人数据', '可能修改了个人设置']
        },
        'post.html': {
            'participants': ['所有用户（浏览）', '已登录用户（交互）'],
            'preconditions': ['用户已访问论坛页面', '帖子存在且可访问'],
            'main_flow': [
                '用户访问帖子详情页面',
                '系统显示帖子标题、内容和元信息',
                '用户浏览帖子内容',
                '用户查看评论和回复',
                '用户浏览相关帖子推荐'
            ],
            'alternative_flow': [
                'A1：发表评论（需登录）',
                '1. 用户在评论框输入评论内容',
                '2. 用户点击"发表评论"按钮',
                '3. 系统检查用户登录状态',
                '4. 如未登录，提示登录或跳转到登录页面',
                '5. 如已登录，系统保存评论并实时显示',
                '6. 系统发送通知给帖子作者（如设置）',
                '',
                'A2：点赞/收藏帖子（需登录）',
                '1. 用户点击"点赞"按钮',
                '2. 系统记录点赞并更新计数',
                '3. 用户点击"收藏"按钮',
                '4. 系统添加帖子到用户收藏',
                '5. 系统显示操作成功提示',
                '',
                'A3：回复评论',
                '1. 用户点击某条评论的"回复"按钮',
                '2. 系统显示回复输入框',
                '3. 用户输入回复内容',
                '4. 用户提交回复',
                '5. 系统保存回复并显示在对应评论下'
            ],
            'postconditions': ['用户浏览了帖子内容', '可能进行了交互操作（评论、点赞等）', '帖子浏览计数增加']
        }
    }
    
    return use_cases.get(filename, {
        'participants': ['用户'],
        'preconditions': ['页面可正常访问'],
        'main_flow': ['用户访问页面', '系统显示页面内容'],
        'alternative_flow': ['无'],
        'postconditions': ['页面访问完成']
    })

def main():
    # 创建文档
    doc = Document()
    doc.core_properties.title = '校务问答机器人系统原型说明'
    doc.core_properties.author = '系统设计团队'
    doc.core_properties.subject = '原型设计与用例说明'
    
    # 设置文档全局字体样式：标题仿宋，正文楷体小四
    set_document_fonts(doc)
    
    # 添加标题
    title = doc.add_heading('校务问答机器人系统', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('原型设计与用例说明文档', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 空行
    
    # 添加文档信息
    info = doc.add_paragraph()
    info.add_run('文档版本：').bold = True
    info.add_run('V1.0')
    
    info = doc.add_paragraph()
    info.add_run('创建日期：').bold = True
    info.add_run('2026年4月19日')
    
    info = doc.add_paragraph()
    info.add_run('文档说明：').bold = True
    info.add_run('本文档包含系统原型页面说明和详细用例文本，用于指导系统开发和测试。')
    
    doc.add_page_break()
    
    # 4.1 原型说明
    doc.add_heading('4.1 原型说明', level=1)
    doc.add_paragraph('本章节详细说明系统各原型页面的功能、目标和输入输出要素。')
    
    # 获取HTML文件，排除final_test.html
    html_files = sorted([f for f in glob.glob('*.html') if os.path.isfile(f) and f != 'final_test.html'])
    
    for filename in html_files:
        with open(filename, 'r', encoding='utf-8') as fh:
            content = fh.read()
        soup = BeautifulSoup(content, 'html.parser')
        title = soup.title.string.strip() if soup.title and soup.title.string else filename
        
        # 获取页面信息
        description = get_page_description(filename, soup)
        goal, outputs = infer_goal_and_outputs(soup, filename)
        inputs = extract_inputs(soup)
        
        # 添加页面标题
        doc.add_heading(f'页面：{filename}', level=2)
        
        # 页面标题和描述
        p_title = doc.add_paragraph()
        p_title.add_run('页面标题：').bold = True
        p_title.add_run(title)
        
        p_desc = doc.add_paragraph()
        p_desc.add_run('页面描述：').bold = True
        p_desc.add_run(description)
        
        # 目标
        p_goal = doc.add_paragraph()
        p_goal.add_run('目标：').bold = True
        p_goal.add_run(goal)
        
        # 输入
        p_in = doc.add_paragraph()
        p_in.add_run('输入：').bold = True
        if inputs:
            for inp in inputs:
                doc.add_paragraph(inp, style='List Bullet')
        else:
            p_in.add_run('无显式表单输入；页面通过导航、点击等交互方式接收用户操作。')
        
        # 输出
        p_out = doc.add_paragraph()
        p_out.add_run('输出：').bold = True
        for out in outputs:
            doc.add_paragraph(out, style='List Bullet')
        
        # 页面功能注释
        doc.add_paragraph()
        note = doc.add_paragraph()
        note.add_run('功能注释：').bold = True
        
        # 根据页面类型添加特定注释
        if filename == 'index.html':
            note.add_run('首页是系统的入口页面，重点突出智能问答功能，同时提供论坛、帮助等功能的快速导航。设计上采用清晰的视觉层次，引导用户使用核心功能。')
        elif filename == 'login.html':
            note.add_run('登录页面采用标签页设计，同时提供登录和注册功能，减少页面跳转。包含密码找回和记住登录状态等辅助功能，提升用户体验。')
        elif filename == 'personal.html':
            note.add_run('个人中心采用侧边栏导航设计，将功能模块化分类。提供个人信息管理、历史记录查看、收藏管理等完整功能，满足用户个性化需求。')
        elif filename == 'forum.html':
            note.add_run('论坛页面设计注重内容组织和发现，提供分类浏览、搜索筛选、热门推荐等多种内容发现方式。鼓励用户参与讨论和内容创建。')
        elif filename == 'post.html':
            note.add_run('帖子详情页面注重阅读体验和交互设计，提供清晰的评论层级、点赞收藏等社交功能，以及相关内容推荐，增加用户粘性。')
        elif filename == 'admin.html':
            note.add_run('管理后台采用专业的管理界面设计，提供数据统计、用户管理、内容审核等完整管理功能。界面设计注重操作效率和数据可视化。')
        elif filename == 'help.html':
            note.add_run('帮助页面采用分类导航设计，提供结构化的帮助内容。包含FAQ、使用指南、联系支持等多种帮助形式，降低用户学习成本。')
        
        doc.add_paragraph()  # 空行分隔
    
    doc.add_page_break()
    
    # 4.2 用例文本
    doc.add_heading('4.2 用例文本', level=1)
    doc.add_paragraph('本章节详细描述系统各功能的使用场景和操作流程，包含参与者、前置条件、主流程、备选流程和后置条件。')
    
    for filename in html_files:
        # 获取用例信息
        use_case = get_use_cases(filename)
        
        # 添加用例标题
        doc.add_heading(f'用例：{filename.replace(".html", "")}功能', level=2)
        
        # 参与者
        p_part = doc.add_paragraph()
        p_part.add_run('参与者：').bold = True
        p_part.add_run('，'.join(use_case['participants']))
        
        # 前置条件
        p_pre = doc.add_paragraph()
        p_pre.add_run('前置条件：').bold = True
        for condition in use_case['preconditions']:
            doc.add_paragraph(condition, style='List Bullet')
        
        # 主流程
        p_main = doc.add_paragraph()
        p_main.add_run('主流程：').bold = True
        for step in use_case['main_flow']:
            doc.add_paragraph(step, style='List Number')
        
        # 备选流程
        p_alt = doc.add_paragraph()
        p_alt.add_run('备选流程：').bold = True
        # 处理备选流程文本（可能包含多行）
        alt_text = '\n'.join(use_case['alternative_flow'])
        for line in alt_text.split('\n'):
            if line.strip():  # 跳过空行
                if line.startswith('A'):  # 子流程标题
                    doc.add_paragraph(line, style='Heading 3')
                else:
                    doc.add_paragraph(line)
        
        # 后置条件
        p_post = doc.add_paragraph()
        p_post.add_run('后置条件：').bold = True
        for condition in use_case['postconditions']:
            doc.add_paragraph(condition, style='List Bullet')
        
        doc.add_paragraph()  # 空行分隔
    
    # 保存文档
    out_path = '原型说明_详细版.docx'
    doc.save(out_path)
    print(f'已生成详细文档：{out_path}')
    
    # 输出统计信息
    print(f'处理页面数量：{len(html_files)}')
    print(f'排除页面：final_test.html')

if __name__ == '__main__':
    main()